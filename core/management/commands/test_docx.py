from django.core.management.base import BaseCommand
from io import BytesIO
import docx
import logging
from django.utils import timezone
from data_ingestion.services.document_processor import DocumentProcessor
from core.services.answer_validation import AnswerValidator
from data_ingestion.models import ProcessingLog
from core.models import Assessment, Course

logger = logging.getLogger(__name__)

class Command(BaseCommand):
    help = 'Test full document processing pipeline including answer validation'

    def add_arguments(self, parser):
        parser.add_argument(
            '--verbose',
            action='store_true',
            help='Show detailed processing output'
        )

    def handle(self, *args, **options):
        verbose = options['verbose']
        
        self.stdout.write("Testing full document processing pipeline...")
        
        try:
            # Create test document
            doc = docx.Document()
            doc.add_paragraph("Question 1: What is inheritance in Java?")
            doc.add_paragraph("Answer: Inheritance allows a class to inherit properties from another class.")
            doc.add_paragraph("Question 2: What is polymorphism?")
            doc.add_paragraph("Answer: Polymorphism allows objects to take many forms.")
            
            # Create test document record
            from django.contrib.auth.models import User
            from data_ingestion.models import DocumentUpload
            
            # Get or create test user
            test_user = User.objects.filter(username='test_user').first()
            if not test_user:
                test_user = User.objects.create_user(
                    username='test_user',
                    password='testpass123'
                )
            
            # Clean up any existing test data
            DocumentUpload.objects.filter(id=999).delete()
            Assessment.objects.filter(assessment_id='test-assessment').delete()
            
            # Create test course and assessment
            test_course = Course.objects.get_or_create(
                course_id='test-course',
                defaults={
                    'title': 'Test Course',
                    'description': 'Course for testing purposes'
                }
            )[0]
            
            test_assessment = Assessment.objects.create(
                assessment_id='test-assessment',
                title='Test Assessment',
                assessment_type='exam',
                course=test_course,
                date=timezone.now(),
                proctored=False
            )
            
            # Create test file in media directory
            import os
            from django.conf import settings
            test_file_path = os.path.join(settings.MEDIA_ROOT, 'test_document.docx')
            doc.save(test_file_path)
            
            # Create test document
            test_doc = DocumentUpload.objects.create(
                id=999,
                file='test_document.docx',
                document_type='docx',
                uploaded_by=test_user,
                status='pending',
                assessment=test_assessment
            )
            
            # Process document
            processor = DocumentProcessor(document_id=test_doc.id)
            validator = AnswerValidator()
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Process the document
            success = processor.process()
            results = {
                'questions': [
                    {
                        'text': q.question_text,
                        'answer': q.student_answer,
                        'validation': {
                            'is_correct': q.is_correct,
                            'score': q.confidence_score,
                            'feedback': q.feedback
                        }
                    }
                    for q in processor.document.questions.all()
                ]
            }
            
            if verbose:
                self.stdout.write("\nProcessing results:")
                for q in results['questions']:
                    self.stdout.write(f"\nQuestion: {q['text']}")
                    self.stdout.write(f"Answer: {q['answer']}")
                    
                    # Validate answer - create proper question object first
                    from core.models import Question
                    question_obj = Question(
                        text=q['text'],
                        question_type='short_answer'
                    )
                    validation = validator.validate_answer(question_obj, q['answer'])
                    self.stdout.write(f"\nDetailed Validation:")
                    self.stdout.write(f"Question: {question_obj.text}")
                    self.stdout.write(f"Answer: {q['answer']}")
                    self.stdout.write(f"Correct: {validation['is_correct']}")
                    self.stdout.write(f"Score: {validation['score']}")
                    self.stdout.write(f"Feedback: {validation['feedback']}")
            
            # Show processing logs
            self.stdout.write("\nProcessing logs:")
            for log in ProcessingLog.objects.order_by('-timestamp')[:5]:
                self.stdout.write(f"{log.timestamp}: {log.message}")
            
            self.stdout.write(self.style.SUCCESS("\nDocument processing test completed"))
            
        except Exception as e:
            logger.exception("Document processing test failed")
            self.stdout.write(self.style.ERROR(f"Test failed: {str(e)}"))
